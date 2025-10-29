#!/usr/bin/env python3
"""
Generate Standardized Bid Proposal Document for LNBL RFP
Professional 1-2 page format with pricing and scope details
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json

# Load calculation data
with open('G:/Shared drives/Energen Ops/2-Sales & Marketing/1-Sales/Bids/LNBL-Whole-Facility-RFP/final-corrected-calculation.json', 'r') as f:
    calc_data = json.load(f)

calc = calc_data['calculation']

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Page 1: BID PROPOSAL SUMMARY
# ============================================================================

# Header/Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BID PROPOSAL')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(54, 96, 146)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Generator Maintenance and Repair Services')
run.font.size = Pt(12)

# RFP Info
doc.add_paragraph()
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Light Grid Accent 1'

info_data = [
    ('RFP Number:', 'ANR-6-2025'),
    ('Customer:', 'Lawrence Berkeley National Laboratory'),
    ('Location:', '1 Cyclotron Road, Berkeley, CA 94720'),
    ('Contract Period:', '1 Year Base + 3 Option Years'),
    ('Proposal Date:', 'October 24, 2025')
]

for idx, (label, value) in enumerate(info_data):
    row = info_table.rows[idx]
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[1].text = value

doc.add_paragraph()

# SCOPE OF SERVICES
heading = doc.add_paragraph()
run = heading.add_run('SCOPE OF SERVICES')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(54, 96, 146)

doc.add_paragraph(
    'Energen Systems proposes to provide comprehensive preventative maintenance and testing services '
    'for thirty-four (34) standby generators at LBNL facilities, as follows:'
)

# Services Table
services_table = doc.add_table(rows=5, cols=4)
services_table.style = 'Light Grid Accent 1'
services_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
header_cells = services_table.rows[0].cells
headers = ['Service', 'Description', 'Frequency', 'Units']
for idx, header_text in enumerate(headers):
    header_cells[idx].text = header_text
    header_cells[idx].paragraphs[0].runs[0].font.bold = True

# Service rows
service_rows = [
    ('Service A', 'Comprehensive Inspection & Testing', 'Annual', '34'),
    ('Service B', 'Oil & Filter Service', 'Annual', '34'),
    ('Service D', 'Fluid Analysis (Oil, Fuel, Coolant)', 'Annual', '34'),
    ('Service E', 'Load Bank Testing (2 Hours)', 'Annual', '34')
]

for idx, (svc, desc, freq, units) in enumerate(service_rows, 1):
    row = services_table.rows[idx]
    row.cells[0].text = svc
    row.cells[1].text = desc
    row.cells[2].text = freq
    row.cells[3].text = units

doc.add_paragraph()

# SERVICE DETAILS SECTION
heading = doc.add_paragraph()
run = heading.add_run('SERVICE DETAILS')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(54, 96, 146)

# Service A Details
p = doc.add_paragraph()
run = p.add_run('Service A - Comprehensive Inspection:')
run.font.bold = True
doc.add_paragraph(
    '• Pre-start checks (battery, fluids, belts, leaks, controls)\n'
    '• Engine running tests (pressures, voltages, frequencies)\n'
    '• Battery service and testing\n'
    '• Belt and hose inspection\n'
    '• Photo documentation (before/after)\n'
    '• Detailed service report\n'
    '• CMMS (Maximo) data entry',
    style='List Bullet'
)

# Service B Details
p = doc.add_paragraph()
run = p.add_run('Service B - Oil & Filter Service:')
run.font.bold = True
doc.add_paragraph(
    '• Drain and replace crankcase oil\n'
    '• Replace oil filters\n'
    '• Replace fuel filters\n'
    '• Coolant condition check and replacement if needed',
    style='List Bullet'
)

# Service D Details
p = doc.add_paragraph()
run = p.add_run('Service D - Fluid Analysis:')
run.font.bold = True
doc.add_paragraph(
    '• Oil sample collection\n'
    '• Fuel sample collection\n'
    '• Coolant sample collection\n'
    '• Laboratory analysis\n'
    '• Written report with recommendations',
    style='List Bullet'
)

# Service E Details
p = doc.add_paragraph()
run = p.add_run('Service E - Load Bank Testing:')
run.font.bold = True
doc.add_paragraph(
    '• 2-hour load bank test at full load\n'
    '• Record voltage, frequency, amperage every 15 minutes\n'
    '• Temperature monitoring\n'
    '• Performance verification\n'
    '• Test completion report',
    style='List Bullet'
)

# Page Break
doc.add_page_break()

# Page 2: PRICING & SPECIAL CONSIDERATIONS
# ============================================================================

heading = doc.add_paragraph()
run = heading.add_run('ANNUAL PRICING SUMMARY')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(54, 96, 146)

# Pricing Table
pricing_table = doc.add_table(rows=6, cols=2)
pricing_table.style = 'Light Grid Accent 1'

labor = float(calc['laborTotal'])
parts = float(calc['partsTotal'])
mobilization = float(calc['mobilizationTotal'])
subtotal = float(calc['subtotal'])

pricing_data = [
    ('Labor', f'${labor:,.2f}'),
    ('Parts & Materials', f'${parts:,.2f}'),
    ('Travel/Mobilization', f'${mobilization:,.2f}'),
    ('', ''),
    ('TOTAL ANNUAL COST', f'${subtotal:,.2f}')
]

for idx, (label, amount) in enumerate(pricing_data):
    row = pricing_table.rows[idx]
    row.cells[0].text = label
    row.cells[1].text = amount
    if idx == 4:  # Total row
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(12)

# Service breakdown
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Pricing by Service:')
run.font.bold = True

breakdown_table = doc.add_table(rows=5, cols=2)
breakdown_table.style = 'Light Grid Accent 1'

services = calc['serviceBreakdown']
service_a = services.get('A - Comprehensive Inspection', {}).get('totalCost', 0)
service_b = services.get('B - Oil & Filter Service', {}).get('totalCost', 0)
service_d = services.get('D - Oil & Fuel Analysis', {}).get('totalCost', 0)
service_e = services.get('E - Load Bank Testing', {}).get('totalCost', 0)

breakdown_data = [
    ('Service A - Comprehensive Inspection', f'${service_a:,.2f}'),
    ('Service B - Oil & Filter Service', f'${service_b:,.2f}'),
    ('Service D - Fluid Analysis', f'${service_d:,.2f}'),
    ('Service E - Load Bank Testing', f'${service_e:,.2f}')
]

for idx, (svc, cost) in enumerate(breakdown_data):
    row = breakdown_table.rows[idx]
    row.cells[0].text = svc
    row.cells[1].text = cost

doc.add_paragraph()

# SPECIAL CONSIDERATIONS
heading = doc.add_paragraph()
run = heading.add_run('SPECIAL CONSIDERATIONS & COST FACTORS')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(54, 96, 146)

doc.add_paragraph(
    'The following site-specific requirements will affect service delivery and have been considered '
    'in our pricing structure:'
)

# Special considerations list
considerations = doc.add_table(rows=6, cols=2)
considerations.style = 'Light List Accent 1'

considerations_data = [
    ('DOE Security Clearance', 'Additional 15-30 minutes per visit for escort/badging procedures'),
    ('Lockout/Tagout', 'Personal lockout required per DOE standards (15-20 minutes per unit)'),
    ('Enhanced Documentation', 'Photo documentation, CMMS entry, detailed reports per unit'),
    ('BAAQMD Compliance', 'Permit compliance verification for permitted engines'),
    ('Prevailing Wage', 'California prevailing wage rates applied per DIR requirements'),
    ('24/7 Emergency Service', 'On-call coverage included; T&M rates apply for emergency calls')
]

for idx, (factor, impact) in enumerate(considerations_data):
    row = considerations.rows[idx]
    row.cells[0].text = factor
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[1].text = impact

doc.add_paragraph()

# COMPLIANCE CERTIFICATIONS
heading = doc.add_paragraph()
run = heading.add_run('COMPLIANCE & CERTIFICATIONS')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(54, 96, 146)

doc.add_paragraph(
    '✓ California C-10 Electrical Contractor License\n'
    '✓ NFPA 70E trained technicians (8 hours minimum)\n'
    '✓ DIR Registration for prevailing wage compliance\n'
    '✓ DOE/LBNL safety standards compliance (10 CFR 851)\n'
    '✓ Insurance: Additional Insured (UC Regents & US Government)\n'
    '✓ BAAQMD permit compliance experience',
    style='List Bullet'
)

doc.add_paragraph()

# PAYMENT TERMS
heading = doc.add_paragraph()
run = heading.add_run('PAYMENT TERMS')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(54, 96, 146)

doc.add_paragraph(
    '• Net 30 Days from invoice date\n'
    '• Tax Exempt (Government Contract)\n'
    '• Invoices submitted to: APINVOICE@lbl.gov\n'
    '• Includes itemized breakdown, parts list, and service reports',
    style='List Bullet'
)

# Footer
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Energen Systems | Generator Maintenance Specialists')
run.font.size = Pt(10)
run.font.italic = True

# Save document
output_file = 'G:/Shared drives/Energen Ops/2-Sales & Marketing/1-Sales/Bids/LNBL-Whole-Facility-RFP/LNBL-Bid-Proposal.docx'
doc.save(output_file)

print("=" * 80)
print("BID PROPOSAL DOCUMENT CREATED")
print("=" * 80)
print(f"\nFile: LNBL-Bid-Proposal.docx")
print(f"\nDocument Structure:")
print(f"  Page 1: Scope of Services & Service Details")
print(f"  Page 2: Pricing Summary & Special Considerations")
print(f"\nTotal Annual Cost: ${subtotal:,.2f}")
print(f"34 Generators | 4 Annual Services per Unit")
print("=" * 80)
